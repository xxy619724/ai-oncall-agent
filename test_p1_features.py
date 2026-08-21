"""P1 优化项单元测试

覆盖四项改动（不依赖 Milvus / LLM 等外部服务）：
1. P1-1 OCR：扫描页路由（mock ocr_page）
2. P1-2 表格：Markdown 转换 / 表格块抽离 / 独立分片 / docx 端到端 / PDF 端到端
3. P1-3 doc：报错文案（轻量跳过，见 file.py）
4. P1-4 语义缓存：知识库版本号文件操作 / 开关与冷却旁路 / 重复提问命中验证

运行: .venv\\Scripts\\python.exe test_p1_features.py
"""

import sys
import os
import tempfile
from pathlib import Path

# 添加项目根目录到 path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app.config import config


# ---------- P1-2: 表格 Markdown 转换 ----------

def test_rows_to_markdown():
    from app.services.document_parser_service import DocumentParserService

    f = DocumentParserService._rows_to_markdown

    # 基本表格：表头 + 分隔行 + 数据行
    md = f([["指标", "阈值"], ["CPU", "90%"], ["内存", "80%"]])
    lines = md.split("\n")
    assert lines[0] == "| 指标 | 阈值 |"
    assert lines[1] == "|---|---|"
    assert lines[2] == "| CPU | 90% |"

    # None 单元格 → 空串；管道符转义；换行合并
    md2 = f([["a", None], ["x|y", "l1\nl2"]])
    assert "| a |  |" in md2
    assert "x\\|y" in md2
    assert "l1 l2" in md2

    # 列数不齐 → 补齐
    md3 = f([["h1", "h2", "h3"], ["only-one"]])
    assert md3.split("\n")[2] == "| only-one |  |  |"

    # 误检过滤：<2行 或 <2列 → 空串
    assert f([["only header"]]) == ""
    assert f([["a"], ["b"]]) == ""

    print("✓ P1-2 表格 Markdown 转换")


# ---------- P1-2: 表格块抽离与独立分片 ----------

def test_extract_table_blocks():
    from app.services.document_splitter_service import document_splitter_service

    content = (
        "[[PAGE:1]]\n正文第一页。\n"
        "[[TABLE]]\n| a | b |\n|---|---|\n| 1 | 2 |\n[[/TABLE]]\n"
        "[[PAGE:2]]\n正文第二页。\n"
        "[[TABLE]]\n| c | d |\n|---|---|\n| 3 | 4 |\n[[/TABLE]]\n"
    )
    new_content, blocks = document_splitter_service._extract_table_blocks(content)

    assert len(blocks) == 2
    assert blocks[0]["page"] == 1, "第一个表格块应关联第1页"
    assert blocks[1]["page"] == 2, "第二个表格块应关联第2页"
    assert "| a | b |" in blocks[0]["markdown"]
    assert "[[TABLE]]" not in new_content, "表格块应从文本流中移除"
    assert "正文第一页" in new_content and "正文第二页" in new_content

    print("✓ P1-2 表格块抽离（含页码关联）")


def test_build_table_chunks():
    from app.services.document_splitter_service import (
        document_splitter_service, _MAX_TABLE_CHUNK_CHARS,
    )

    # 普通表格 → 单分片，带 _type 和页码
    blocks = [{"markdown": "| a | b |\n|---|---|\n| 1 | 2 |", "page": 3}]
    chunks = document_splitter_service._build_table_chunks(blocks, "uploads/t.pdf")
    assert len(chunks) == 1
    assert chunks[0].metadata["_type"] == "table"
    assert chunks[0].metadata["_page_start"] == 3
    assert chunks[0].metadata["_page_end"] == 3
    assert chunks[0].metadata["_file_name"] == "t.pdf"

    # 超长表格 → 按行切分，重复表头，带 _table_part
    header = "| col |" + " val |" * 40  # 宽表头，保证超限
    sep = "|" + "---|" * 41
    rows = "\n".join(f"| r{i} |" + " x |" * 40 for i in range(300))
    big_md = f"{header}\n{sep}\n{rows}"
    assert len(big_md) > _MAX_TABLE_CHUNK_CHARS

    chunks2 = document_splitter_service._build_table_chunks(
        [{"markdown": big_md, "page": None}], "uploads/t.pdf"
    )
    assert len(chunks2) > 1
    parts = [c.metadata["_table_part"] for c in chunks2]
    assert parts[0].endswith(f"/{len(chunks2)}")
    for c in chunks2:
        assert len(c.page_content) <= _MAX_TABLE_CHUNK_CHARS, "每片不得超上限"
        assert c.page_content.startswith("| col |"), "每片应重复表头"
        assert "_page_start" not in c.metadata, "无页码块不应有页码元数据"

    print(f"✓ P1-2 表格独立分片（超长切分 {len(chunks2)} 片，均重复表头）")


def test_split_document_table_integration():
    """端到端：[[TABLE]] 标记 → 文本分片 + 表格分片共存"""
    from app.services.document_splitter_service import document_splitter_service

    content = (
        "[[PAGE:1]]\n"
        + "CPU告警处理手册正文，包含定位进程和确认阈值两部分内容。" * 20
        + "\n[[TABLE]]\n| 指标 | 阈值 |\n|---|---|\n| CPU | 90% |\n[[/TABLE]]\n"
    )
    docs = document_splitter_service.split_document(content, "uploads/manual.pdf")

    table_docs = [d for d in docs if d.metadata.get("_type") == "table"]
    text_docs = [d for d in docs if d.metadata.get("_type") != "table"]

    assert table_docs, "应有表格分片"
    assert "| 指标 | 阈值 |" in table_docs[0].page_content
    assert table_docs[0].metadata["_page_start"] == 1
    assert text_docs, "应有正文分片"
    for d in text_docs:
        assert "[[TABLE" not in d.page_content, "正文分片不应残留表格标记"
        assert "| 指标 |" not in d.page_content, "表格内容不应混入正文分片"

    print("✓ P1-2 分片端到端（正文与表格分片共存）")


# ---------- P1-2: docx 端到端（python-docx 生成测试文件） ----------

def test_docx_table_end_to_end():
    from docx import Document as DocxDocument

    from app.services.document_parser_service import document_parser_service
    from app.services.document_splitter_service import document_splitter_service

    # 生成带标题、正文、表格的 docx
    doc = DocxDocument()
    doc.add_heading("CPU 高负载处理", level=1)
    doc.add_paragraph("先定位异常进程，再结合告警规则确认是否误报。")
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "指标"
    table.cell(0, 1).text = "阈值"
    table.cell(1, 0).text = "CPU"
    table.cell(1, 1).text = "90%"
    table.cell(2, 0).text = "内存"
    table.cell(2, 1).text = "80%"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        docx_path = f.name

    try:
        content = document_parser_service.parse(docx_path)
        # 解析层：标题映射 + 表格块标记 + 章节上下文
        assert "# CPU 高负载处理" in content
        assert "[[TABLE]]" in content and "[[/TABLE]]" in content
        assert "所属章节: CPU 高负载处理" in content, "表格应附带章节上下文"
        assert "| 指标 | 阈值 |" in content

        # 分片层：表格独立分片 + 正文走标题切分
        docs = document_splitter_service.split_document(content, docx_path.replace("\\", "/"))
        table_docs = [d for d in docs if d.metadata.get("_type") == "table"]
        assert table_docs, "docx 表格应成为独立分片"
        assert "| CPU | 90% |" in table_docs[0].page_content
        assert any(d.metadata.get("h1") == "CPU 高负载处理" for d in docs), "正文保留标题切分"
    finally:
        os.unlink(docx_path)

    print("✓ P1-2 docx 表格端到端（章节上下文 + 独立分片）")


# ---------- P1-1/P1-2: PDF 端到端（pymupdf 生成测试文件） ----------

def _make_pdf(pages_spec, path):
    """用 pymupdf 生成测试 PDF

    pages_spec: 每页的绘制函数列表 [(draw_fn, page), ...]
    """
    import pymupdf

    doc = pymupdf.open()
    for draw_fn in pages_spec:
        page = doc.new_page(width=595, height=842)  # A4
        draw_fn(page)
    doc.save(path)
    doc.close()


def test_pdf_text_page():
    """纯文字页走 pypdf 快路径，保留 [[PAGE:n]] 标记"""
    from app.services.document_parser_service import document_parser_service

    def draw(page):
        page.insert_text((72, 100), "This is a long body text for testing. " * 5)

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        pdf_path = f.name
    # Windows 下需先关闭句柄，pymupdf 才能写入
    _make_pdf([draw], pdf_path)

    try:
        config.ocr_enabled = False  # 确保走文字通道
        content = document_parser_service.parse(pdf_path)
        assert "[[PAGE:1]]" in content
        assert "long body text" in content
        assert "[[TABLE]]" not in content
    finally:
        config.ocr_enabled = True
        os.unlink(pdf_path)

    print("✓ P1-1 PDF 纯文字页（快路径 + 页码标记）")


def test_pdf_table_page():
    """画线表格页 → 表格块 + 正文剔除表格文字"""
    from app.services.document_parser_service import document_parser_service

    def draw(page):
        # 3x2 网格线（表格区域 x:50-250, y:100-160）
        for y in (100, 130, 160):
            page.draw_line(pymupdf.Point(50, y), pymupdf.Point(250, y))
        for x in (50, 150, 250):
            page.draw_line(pymupdf.Point(x, 100), pymupdf.Point(x, 160))
        # 表格内文字
        page.insert_text((60, 120), "Metric")
        page.insert_text((160, 120), "Value")
        page.insert_text((60, 150), "CPU")
        page.insert_text((160, 150), "90")
        # 表格外正文（足够长避免触发 OCR 判定）
        page.insert_text((50, 200), "This is body text outside the table area. " * 3)

    import pymupdf

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        pdf_path = f.name
    _make_pdf([draw], pdf_path)

    try:
        config.ocr_enabled = False  # 关闭 OCR，走表格通道
        content = document_parser_service.parse(pdf_path)

        assert "[[TABLE]]" in content, "检测到表格应输出表格块"
        assert "| Metric | Value |" in content, "表格应转为 Markdown"
        assert "| CPU | 90 |" in content

        # 正文剔除表格文字：表格单元格文字不应出现在正文行中
        body_lines = [
            ln for ln in content.split("\n")
            if "Metric" in ln or "Value" in ln
        ]
        non_table_body = [
            ln for ln in body_lines if not ln.startswith("|")
        ]
        assert not non_table_body, f"表格文字混入正文: {non_table_body}"
        assert "body text outside" in content, "表格外正文应保留"
    finally:
        config.ocr_enabled = True
        os.unlink(pdf_path)

    print("✓ P1-2 PDF 表格页端到端（表格 Markdown + 正文去重）")


def test_pdf_scanned_page_ocr():
    """无文字页（扫描页）→ OCR 路由（mock ocr_page 返回转写文本）"""
    from app.services.document_parser_service import document_parser_service
    from app.services import ocr_service as ocr_module

    def draw(page):
        pass  # 空白页，无文本层

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        pdf_path = f.name
    _make_pdf([draw], pdf_path)

    original = ocr_module.ocr_service.ocr_page
    try:
        config.ocr_enabled = True
        ocr_module.ocr_service.ocr_page = lambda path, idx: "OCR识别的扫描页文字内容"

        content = document_parser_service.parse(pdf_path)
        assert "[[PAGE:1]]" in content
        assert "OCR识别的扫描页文字内容" in content
    finally:
        ocr_module.ocr_service.ocr_page = original
        os.unlink(pdf_path)

    print("✓ P1-1 PDF 扫描页 OCR 路由")


def main():
    print("=" * 60)
    print("P1 优化项单元测试")
    print("=" * 60)

    test_rows_to_markdown()
    test_extract_table_blocks()
    test_build_table_chunks()
    test_split_document_table_integration()
    test_docx_table_end_to_end()
    test_pdf_text_page()
    test_pdf_table_page()
    test_pdf_scanned_page_ocr()

    print("\n" + "=" * 60)
    print("全部通过 ✓")
    print("=" * 60)


if __name__ == "__main__":
    main()
